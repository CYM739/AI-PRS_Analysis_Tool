# src/logic/data_processing.py
import os
import pandas as pd
import numpy as np
import statsmodels.formula.api as smf
from docx import Document
import io
from datetime import datetime
from zipfile import ZipFile
from .models import OLSWrapper, SVRWrapper, RandomForestWrapper
from .helpers import _add_polynomial_terms

def analyze_csv(full_dataframe):
    """
    Inspects the input dataframe to prepare it for analysis.
    This function intelligently:
    1.  Detects if the first row contains variable descriptions.
    2.  Identifies independent vs. dependent variables based on naming conventions.
    3.  Cleans numeric columns that may have been read as strings due to formatting (e.g., commas).
    4.  Calculates key statistics (min, max, second min) for each variable.
    5.  Extracts all unique values for use in dropdowns.
    6.  Automatically detects which independent variables are binary (contain only 0s and 1s).
    """
    first_row = full_dataframe.iloc[0]
    is_description_row = any(isinstance(item, str) for item in first_row)

    all_vars = full_dataframe.columns.tolist()

    # Separate data from descriptions if the first row is descriptive text.
    if is_description_row:
        descriptions_row = first_row
        data_df = full_dataframe.iloc[1:].reset_index(drop=True)
        variable_descriptions = {var: descriptions_row[var] for var in all_vars}
    else:
        data_df = full_dataframe.copy()
        variable_descriptions = {var: var for var in all_vars}

    for col in data_df.columns:
        if data_df[col].dtype == 'object':
            try:
                cleaned_col = data_df[col].str.replace(',', '', regex=False)
                data_df[col] = pd.to_numeric(cleaned_col)
            except (AttributeError, ValueError, TypeError):
                pass

    ignore_prefixes = ["Cell_", "Zebrafish_", "Mouse_", "Patient_", "Control_"]
    independent_vars = [
        col for col in all_vars
        if col[0].isalpha() and not any(col.startswith(prefix) for prefix in ignore_prefixes)
    ]
    dependent_vars = [
        col for col in all_vars
        if any(col.startswith(prefix) for prefix in ignore_prefixes)
    ]

    variable_stats = {}
    unique_variable_values = {}
    detected_binary_vars = []

    for var in all_vars:
        numeric_col = pd.to_numeric(data_df[var], errors='coerce').dropna()
        if not numeric_col.empty:
            unique_sorted = np.sort(numeric_col.unique())
            unique_list = unique_sorted.tolist()
            unique_variable_values[var] = unique_list
            min_val = unique_sorted[0]
            max_val = unique_sorted[-1]
            second_min_val = unique_sorted[1] if len(unique_sorted) > 1 else min_val
            variable_stats[var] = (min_val, second_min_val, max_val)

            if var in independent_vars:
                if set(unique_list) == {0.0, 1.0} or set(unique_list) == {0, 1}:
                    detected_binary_vars.append(var)

    special_values_map = {}
    for col_name in all_vars:
        numeric_series = pd.to_numeric(data_df[col_name], errors='coerce')
        non_numeric_mask = numeric_series.isna() & data_df[col_name].notna()
        if non_numeric_mask.any():
            special_values = data_df[col_name][non_numeric_mask].unique().tolist()
            special_values_map[col_name] = special_values

    return (data_df, all_vars, independent_vars, dependent_vars,
            variable_stats, special_values_map, unique_variable_values,
            variable_descriptions, detected_binary_vars)

def expand_terms(dataframe, all_alphabet_vars):
    """A convenient wrapper for the polynomial term generation function."""
    _add_polynomial_terms(dataframe, all_alphabet_vars)

def generate_model_formula(C_col, all_alphabet_vars):
    """
    Constructs the model formula string required by the `statsmodels` library.
    """
    terms = []
    for header in all_alphabet_vars:
        terms.extend([header, f"{header}_sq"])

    if len(all_alphabet_vars) > 1:
        interaction_terms = [f"{var1}*{var2}" for i, var1 in enumerate(all_alphabet_vars) for var2 in all_alphabet_vars[i + 1:]]
        terms.extend(interaction_terms)

    return f"{C_col} ~ " + " + ".join(terms)

def run_analysis(dataframe, independent_vars, dependent_var, model_type, model_params={}, variable_descriptions=None):
    """
    Main analysis engine. Fits the specified model and returns a wrapped model object.
    """
    if model_type == 'Polynomial OLS':
        model_formula = generate_model_formula(dependent_var, independent_vars)
        ols_model = smf.ols(model_formula, data=dataframe).fit()
        return OLSWrapper(ols_model, model_formula)

    elif model_type == 'SVR':
        svr_wrapper = SVRWrapper(independent_vars)
        svr_wrapper.fit(dataframe, dependent_var, **model_params)
        return svr_wrapper

    elif model_type == 'Random Forest':
        rf_wrapper = RandomForestWrapper(independent_vars)
        # Pass variable_descriptions to the fit method
        model_params['variable_descriptions'] = variable_descriptions
        rf_wrapper.fit(dataframe, dependent_var, **model_params)
        return rf_wrapper
        
    else:
        raise ValueError(f"Unknown model type: {model_type}")

def predict_surface(model, all_alphabet_vars, x_var, y_var, fixed_vars_dict, x_range, y_range):
    x_values = np.linspace(x_range[0], x_range[1], 50)
    y_values = np.linspace(y_range[0], y_range[1], 50)
    x_grid, y_grid = np.meshgrid(x_values, y_values)
    predict_df = pd.DataFrame({x_var: x_grid.flatten(), y_var: y_grid.flatten()})

    for var, val in fixed_vars_dict.items():
        predict_df[var] = val
    predicted_z = model.predict(predict_df)

    if isinstance(predicted_z, pd.Series):
        predicted_z = predicted_z.values

    z_grid = predicted_z.reshape(x_grid.shape)
    return x_grid, y_grid, z_grid

def generate_optimization_report(report_data, variable_descriptions):
    """
    Generates a professional, well-formatted .docx report from the results of an
    optimization run. The report includes model details, settings, and final results.
    """
    doc = Document()
    doc.add_heading('Optimization Run Report', 0)

    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Optimization Type: {report_data.get('optimization_type', 'N/A')}")

    doc.add_heading('Model Information', level=1)
    doc.add_paragraph(f"Model Name(s): {report_data.get('model_name', 'N/A')}")
    formula = report_data.get('model_formula', 'N/A')
    doc.add_paragraph("Model Formula:")
    doc.add_paragraph(str(formula), style='Intense Quote')

    doc.add_heading('Model Coefficients', level=2)
    params = report_data.get('model_params')
    if params is not None:
        if isinstance(params, dict) and all(isinstance(p, pd.Series) for p in params.values()):
            for model_title, model_params in params.items():
                doc.add_paragraph(f"Coefficients for: {model_title}", style='Heading 3')
                table = doc.add_table(rows=1, cols=2, style='Table Grid')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Term'
                hdr_cells[1].text = 'Coefficient'
                for term, coeff in model_params.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = term
                    row_cells[1].text = f"{coeff:.6f}"
                doc.add_paragraph()
        elif isinstance(params, pd.Series):
            table = doc.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Term'
            hdr_cells[1].text = 'Coefficient'
            for term, coeff in params.items():
                row_cells = table.add_row().cells
                row_cells[0].text = term
                row_cells[1].text = f"{coeff:.6f}"
    else:
        doc.add_paragraph("Coefficient data was not available.")

    doc.add_heading('Optimization Settings', level=1)
    settings = report_data.get('settings', {})
    for key, value in settings.items():
        if key.lower() != 'bounds':
            doc.add_paragraph(f"{key}: {value}")

    doc.add_paragraph("Variable Bounds:")
    bounds = settings.get('Bounds', {})
    if bounds:
        table = doc.add_table(rows=1, cols=3, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Variable'
        hdr_cells[1].text = 'Min Bound'
        hdr_cells[2].text = 'Max Bound'
        for var, (min_b, max_b) in bounds.items():
            row_cells = table.add_row().cells
            row_cells[0].text = variable_descriptions.get(var, var)
            row_cells[1].text = f"{min_b:.4f}"
            row_cells[2].text = f"{max_b:.4f}"

    doc.add_heading('Results', level=1)
    results = report_data.get('results', {})
    doc.add_paragraph(f"Status: {results.get('Status', 'N/A')}")

    if 'Top Solutions' in results and isinstance(results['Top Solutions'], list):
        for result_item in results['Top Solutions']:
            doc.add_heading(f"Rank {result_item['Rank']} Solution", level=2)
            doc.add_paragraph(f"Objective Outcome: {result_item['Final Objective Outcome']:.4f}")
            doc.add_paragraph(f"Constraint Outcome: {result_item['Final Constraint Outcome']:.4f}")
            doc.add_paragraph("Optimal Dosages:")
            dosages = result_item.get('Optimal Dosages', {})
            if dosages:
                table = doc.add_table(rows=1, cols=2, style='Table Grid')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Variable'
                hdr_cells[1].text = 'Optimal Value'
                for var, val in dosages.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = variable_descriptions.get(var, var)
                    row_cells[1].text = f"{val:.4f}"
                doc.add_paragraph()
    else:
        if 'Final Outcome' in results:
            doc.add_paragraph(f"Final Outcome: {results['Final Outcome']:.4f}")
        if 'Final Objective Outcome' in results:
            doc.add_paragraph(f"Final Objective Outcome: {results['Final Objective Outcome']:.4f}")
            doc.add_paragraph(f"Final Constraint Outcome: {results['Final Constraint Outcome']:.4f}")
        if 'Final Outcomes' in results:
            doc.add_paragraph("Final Outcomes:")
            for model, outcome in results['Final Outcomes'].items():
                doc.add_paragraph(f"  - {variable_descriptions.get(model, model)}: {outcome:.4f}", style='List Bullet')

        doc.add_paragraph("Optimal Dosages:")
        dosages = results.get('Optimal Dosages', {})
        if dosages:
            table = doc.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Variable'
            hdr_cells[1].text = 'Optimal Value'
            for var, val in dosages.items():
                row_cells = table.add_row().cells
                row_cells[0].text = variable_descriptions.get(var, var)
                row_cells[1].text = f"{val:.4f}"

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream.getvalue()

def calculate_synergy(dataframe, drug1_name, drug2_name, effect_name, model='gamma'):
    """
    Dispatcher function to calculate drug synergy using a specified model.
    """
    response_matrix = dataframe.pivot_table(
        index=drug1_name, columns=drug2_name, values=effect_name
    )

    if model.lower() == 'gamma':
        expected_matrix = _calculate_hsa_expected(response_matrix)
        synergy_matrix = response_matrix - expected_matrix
        return synergy_matrix
    elif model.lower() == 'hsa':
        expected_matrix = _calculate_hsa_expected(response_matrix)
        synergy_matrix = response_matrix - expected_matrix
        return synergy_matrix
    elif model.lower() == 'loewe':
        return _calculate_loewe_placeholder(response_matrix)
    else:
        raise ValueError(f"Unknown synergy model: {model}")

def _calculate_hsa_expected(response_matrix):
    """
    Calculates the expected response matrix based on the Highest Single Agent (HSA) model.
    """
    drug1_effects = response_matrix.loc[:, 0]
    drug2_effects = response_matrix.loc[0, :]

    expected_matrix = pd.DataFrame(index=response_matrix.index, columns=response_matrix.columns)
    for d1_dose in response_matrix.index:
        for d2_dose in response_matrix.columns:
            expected_matrix.loc[d1_dose, d2_dose] = max(drug1_effects[d1_dose], drug2_effects[d2_dose])

    return expected_matrix

def _calculate_loewe_placeholder(response_matrix):
    """
    Placeholder for the Loewe Additivity model.
    """
    synergy_matrix = response_matrix - response_matrix
    return synergy_matrix

def generate_combined_predictions_csv(wrapped_models, data_df, independent_vars):
    """
    Generates a single CSV file content with actual and predicted values for all models.
    """
    all_predictions_df = pd.DataFrame(index=data_df.index)

    for model_name, model_obj in wrapped_models.items():
        # Prepare data for prediction
        clean_df = data_df.dropna(subset=[model_name] + independent_vars).copy()
        y_actual = clean_df[model_name]
        y_predicted = model_obj.predict(clean_df)

        # Add to the combined dataframe
        all_predictions_df[f'Actual_{model_name}'] = y_actual
        all_predictions_df[f'Predicted_{model_name}'] = pd.Series(y_predicted, index=y_actual.index)

    csv_buffer = io.StringIO()
    all_predictions_df.to_csv(csv_buffer, index_label='Experiment_Index')
    return csv_buffer.getvalue().encode('utf-8')

def generate_combined_report(report_data_list, variable_descriptions):
    """
    Generates a single .docx report containing the results from multiple models.
    """
    doc = Document()
    doc.add_heading('Combined Analysis Report', 0)
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for report_data in report_data_list:
        model_name = report_data.get('model_name', 'N/A')
        doc.add_heading(f"Results for: {model_name}", level=1)

        # You can reuse parts of the single report generation here
        # For simplicity, we'll add the model summary
        formula = report_data.get('model_formula', 'N/A')
        doc.add_paragraph("Model Formula/Type:")
        doc.add_paragraph(str(formula), style='Intense Quote')

        doc.add_heading('Model Summary', level=2)
        # Assuming the summary is pre-formatted text
        summary_text = report_data.get('model_summary', 'Summary not available.')
        doc.add_paragraph(summary_text)
        doc.add_page_break()

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()